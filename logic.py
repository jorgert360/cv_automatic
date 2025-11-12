# logic.py (Versi√≥n con manejo de error de L√≠mite de Tarifa)

import os
import re
import json
import google.generativeai as genai
from pdfminer.high_level import extract_text
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google.api_core import exceptions as google_exceptions # <-- 1. IMPORTAR EXCEPCIONES DE GOOGLE

# --- INICIO DE LA MODIFICACI√ìN (Error Personalizado) ---
# 2. Creamos un error personalizado que nuestra app pueda entender
class RateLimitError(Exception):
    pass
# --- FIN DE LA MODIFICACI√ìN ---

def limpiar_texto_para_xml(texto):
    if texto is None: return ""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', texto)

def configurar_gemini():
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        raise ValueError("No se encontr√≥ la variable de entorno GEMINI_API_KEY.")
    genai.configure(api_key=api_key)

def extraer_texto_pdf(ruta_pdf):
    try:
        texto_extraido = extract_text(ruta_pdf)
        return limpiar_texto_para_xml(texto_extraido)
    except Exception as e:
        raise RuntimeError(f"Error al leer el PDF: {e}")

def extraer_texto_docx(ruta_docx):
    try:
        doc = Document(ruta_docx)
        full_text = [para.text for para in doc.paragraphs]
        return limpiar_texto_para_xml('\n'.join(full_text))
    except Exception as e:
        raise RuntimeError(f"Error al leer el DOCX: {e}")

def analizar_y_optimizar_con_gemini(texto_cv, texto_oferta):
    print("ü§ñ Analizando y optimizando el CV con IA...")
    
    model = genai.GenerativeModel("models/gemini-2.5-pro")
    print(f"DEBUG: Intentando usar el modelo: {model.model_name}")

    prompt = f"""
    Act√∫a como una coach de carrera de √©lite y experta en reclutamiento C-Suite, alineada con las 
    filosof√≠as de expertos como Andrew LaCivita y EdnaJobs. Tu objetivo es transformar un CV 
    para que no solo supere el ATS, sino que impresione al reclutador humano. Tu enfoque es la 
    "contrataci√≥n basada en habilidades" y la "cuantificaci√≥n del impacto".

    CV ORIGINAL: --- {texto_cv} ---
    OFERTA DE TRABAJO: --- {texto_oferta} ---

    TAREAS:

    1.  **Optimiza el CV (Formato JSON):** (Esta tarea no cambia) Reestructura el contenido del CV en el formato 
        JSON de salida.
        a.  **Perfil Profesional:** Re-escribe un "Perfil Profesional" de alto impacto (3-4 l√≠neas) 
            que act√∫e como un "gancho" y sea un "espejo" de la OFERTA ESPEC√çFICA.
        b.  **Experiencia Profesional:** Adapta los logros. Reemplaza el lenguaje pasivo 
            (ej: "responsable de") por **verbos de acci√≥n potentes**. Donde sea posible, 
            **CUANTIFICA** el impacto usando el m√©todo de las "8 Grandes".
        c.  **Coherencia ATS:** Aseg√∫rate de que las palabras clave cr√≠ticas de la OFERTA DE TRABAJO 
            se reflejen en el perfil y la experiencia.

    2.  **Genera Retroalimentaci√≥n Estrat√©gica (Lista de strings):** Crea un an√°lisis en 
        dos partes para el candidato.
        a.  **Paso 1: Fortalezas Clave (El primer √≠tem de la lista):** Comienza la retroalimentaci√≥n 
            con un p√°rrafo positivo y alentador. Identifica las 2-3 **fortalezas y habilidades** principales del candidato que S√ç se alinean perfectamente con la oferta de trabajo.
        b.  **Paso 2: Consejos Accionables (Los siguientes √≠tems):** Despu√©s del inicio positivo, 
            contin√∫a con 3-4 consejos accionables (An√°lisis de Brecha Cr√≠tica, 
            Oportunidad de Impacto, Movimiento Estrat√©gico).

    RESPUESTA: Devuelve tu respuesta √∫nicamente en formato JSON. Aseg√∫rate de que los campos que 
    son listas (como experiencia_profesional, educacion, idiomas, retroalimentacion) sean siempre 
    listas, incluso si est√°n vac√≠as ([]), nunca nulos. La estructura debe ser:
    {{
      "cv_optimizado": {{
        "nombre": "Nombre Completo",
        "contacto": {{ "email": "", "telefono": "", "linkedin": "", "ciudad": "" }},
        "perfil_profesional": "Perfil reescrito.",
        "experiencia_profesional": [ {{ "cargo": "", "empresa": "", "ciudad": "", "periodo": "", "logros": ["Logro 1.", "Logro 2."] }} ],
        "educacion": [ {{ "titulo": "", "institucion": "", "periodo": "" }} ],
        "habilidades": {{ "tecnicas": ["Habilidad 1"], "competencias": ["Competencia 1"] }},
        "idiomas": [ {{ "idioma": "Idioma", "nivel": "Nivel" }} ]
      }},
      "retroalimentacion": [
        "**Tus Fortalezas Clave:** Eres un candidato fuerte para este rol gracias a tu experiencia en [Habilidad 1] y [Habilidad 2].",
        "**An√°lisis de Brecha Cr√≠tica:** Consejo 1.",
        "**Oportunidad de Impacto:** Consejo 2.",
        "**Movimiento Estrat√©gico:** Consejo 3."
      ]
    }}
    """
    
    # --- INICIO DE LA MODIFICACI√ìN (Atrapar Error 429) ---
    # 3. A√±adimos un try/except m√°s espec√≠fico
    try:
        response = model.generate_content(prompt)
        if not response.parts:
            raise RuntimeError("La respuesta de la IA fue bloqueada, posiblemente por pol√≠ticas de seguridad.")
        json_text = response.text.strip().replace("```json", "").replace("```", "")
        return json.loads(json_text)
    
    except google_exceptions.ResourceExhausted as e:
        # ¬°Este es el error de "demasiadas solicitudes" (429)!
        print(f"‚ö†Ô∏è Error de L√≠mite de Tasa de API de Gemini (429): {e}")
        # Lanzamos nuestro error personalizado para que app.py lo atrape
        raise RateLimitError("API de Gemini sobrecargada. Por favor, int√©ntelo de nuevo en un minuto.")
    
    except Exception as e:
        # Errores generales (bloqueo de seguridad, etc.)
        raise RuntimeError(f"Error al procesar la respuesta de Gemini: {e}")
    # --- FIN DE LA MODIFICACI√ìN ---

def crear_docx_optimizado(ruta_completa_salida, data):
    # ... (Esta funci√≥n permanece exactamente igual) ...
    print(f"üé® Creando documento Word en: {ruta_completa_salida}")
    if not data or 'cv_optimizado' not in data:
        raise ValueError("No se recibieron datos v√°lidos para crear el documento.")
    
    cv = data.get('cv_optimizado') or {}
    retro = data.get('retroalimentacion') or []
    
    doc = Document()
    style = doc.styles['Normal']; font = style.font; font.name = 'Calibri'; font.size = Pt(11)
    
    doc.add_paragraph(cv.get('nombre', '')).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.size = Pt(20); doc.paragraphs[-1].runs[0].font.bold = True
    
    contacto = cv.get('contacto') or {}
    contact_items = [contacto.get('email'), contacto.get('telefono'), contacto.get('linkedin'), contacto.get('ciudad')]
    contact_line = " | ".join(filter(None, contact_items))
    doc.add_paragraph(contact_line).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].runs[0].font.size = Pt(10)
    
    doc.add_heading('Perfil Profesional', level=1); doc.add_paragraph(cv.get('perfil_profesional', ''))
    
    doc.add_heading('Experiencia Profesional', level=1)
    for exp in cv.get('experiencia_profesional') or []:
        p_cargo = doc.add_paragraph(); p_cargo.add_run(exp.get('cargo', '')).bold = True
        empresa_line = f"{exp.get('empresa', '')} | {exp.get('ciudad', '')} | {exp.get('periodo', '')}"
        p_empresa = doc.add_paragraph(); p_empresa.add_run(empresa_line).italic = True
        p_empresa.paragraph_format.space_before = Pt(0); p_empresa.paragraph_format.space_after = Pt(4)
        for logro in exp.get('logros') or []:
            doc.add_paragraph(logro, style='List Bullet')
        doc.add_paragraph()
        
    doc.add_heading('Educaci√≥n', level=1)
    for edu in cv.get('educacion') or []:
        p_edu = doc.add_paragraph(); p_edu.add_run(edu.get('titulo', '')).bold = True
        p_edu.add_run(f"\n{edu.get('institucion', '')} | {edu.get('periodo', '')}")
    
    doc.add_heading('Habilidades', level=1)
    habilidades = cv.get('habilidades') or {}
    for key, title in habilidades.items():
        if title and isinstance(title, list):
            p_hab = doc.add_paragraph()
            p_hab.add_run(key.replace('_', ' ').replace('-', ' ').title() + ': ').bold = True
            p_hab.add_run(", ".join(title))

    doc.add_heading('Idiomas', level=1)
    for idioma_info in cv.get('idiomas') or []:
        if isinstance(idioma_info, dict):
            texto_idioma = f"{idioma_info.get('idioma', '')}: {idioma_info.get('nivel', '')}"
            doc.add_paragraph(texto_idioma)
        else:
            doc.add_paragraph(idioma_info)

    doc.save(ruta_completa_salida)
    print("‚úÖ Documento guardado correctamente.")

def procesar_cv_completo(ruta_pdf_cv, texto_oferta, output_folder):
    configurar_gemini()
    texto_cv = extraer_texto_pdf(ruta_pdf_cv)

    MAX_CARACTERES_CV = 15000 
    
    if len(texto_cv) > MAX_CARACTERES_CV:
        print(f"‚ö†Ô∏è Alerta: El CV era muy largo ({len(texto_cv)} caracteres). Se ha truncado a {MAX_CARACTERES_CV}.")
        texto_cv = texto_cv[:MAX_CARACTERES_CV]

    datos_optimizados = analizar_y_optimizar_con_gemini(texto_cv, texto_oferta)
    
    if datos_optimizados:
        nombre_archivo_salida = "cv_optimizado.docx"
        ruta_completa_salida = os.path.join(output_folder, nombre_archivo_salida)
        crear_docx_optimizado(ruta_completa_salida, datos_optimizados)
        
        retroalimentacion = datos_optimizados.get('retroalimentacion', [])
        nombre_candidato = datos_optimizados.get('cv_optimizado', {}).get('nombre', 'Candidato')
        
        return nombre_archivo_salida, retroalimentacion, nombre_candidato
    else:
        return None, None, None